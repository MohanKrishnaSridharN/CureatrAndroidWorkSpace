import webbrowser
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import TimeoutException
import selenium.webdriver.support.expected_conditions as EC
from selenium.webdriver.support import expected_conditions as EC
import selenium.webdriver.support.ui as ui
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import *
import Config
from Config import *
import Constants
from Constants import *
import XlsxReader
from XlsxReader import *
import time
import os,sys
import os.path
from DriverScript import logger,handler
import logging
import logging.handlers
import PIL
from PIL import ImageChops
from PIL import Image
from itertools import izip
from selenium.webdriver.common.keys import Keys
import BackEndDrivers
import datetime
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from pyPdf import PdfFileReader
from docx import Document
import unittest
from appium import webdriver

def LaunchWebBrowser(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if browser=="Android-6.0&5.1":
			if user=="user1":
				desired_caps = {}
				desired_caps['platformName'] = 'Android'
				desired_caps['platformVersion'] = '6.0'
				desired_caps['deviceName'] = 'Moto'
				desired_caps['app'] = '/Users/paradigm/dev/cureatr_android/Cureatr/build/outputs/apk/Cureatr-dev-debug.apk'
				driver = webdriver.Remote('http://172.16.93.1:4725/wd/hub', desired_caps)
				#time.sleep(10)
				#element=driver.find_element_by_xpath("//android.widget.CheckedTextView[@text='DaVita']")
				#element.click()
				return driver, "PASS"
			else:
				desired_caps = {}
				desired_caps['platformName'] = 'Android'
				desired_caps['platformVersion'] = '5.1'
				desired_caps['deviceName'] = 'Moto'
				desired_caps['app'] = '/Users/paradigm/dev/cureatr_android/Cureatr/build/outputs/apk/Cureatr-dev-debug.apk'
				driver = webdriver.Remote('http://172.16.93.1:4726/wd/hub', desired_caps)
				#time.sleep(10)
				#element=driver.find_element_by_xpath("//android.widget.CheckedTextView[@text='CAIPA']")
				#element.click()
				return driver, "PASS"
		elif browser=="Chrome":
			if user=="user1":
				desired_caps = {}
				desired_caps['platformName'] = 'Android'
				desired_caps['platformVersion'] = '6.0'
				desired_caps['deviceName'] = 'Moto'
				desired_caps['app'] = '/Users/macmini/dev/CureatrAndroidWorkSpace/APK/CureatrPlay.apk'
				driver = webdriver.Remote('http://192.168.73.1:4725/wd/hub', desired_caps)
				time.sleep(10)
				element=driver.find_element_by_xpath("//android.widget.CheckedTextView[@text='DaVita']")
				element.click()
				return driver, "PASS"
			else:
				desired_caps = {}
				desired_caps['platformName'] = 'Android'
				desired_caps['platformVersion'] = '4.4'
				desired_caps['deviceName'] = 'Moto'
				desired_caps['app'] = '/Users/macmini/dev/CureatrAndroidWorkSpace/APK/CureatrPlay.apk'
				driver = webdriver.Remote('http://192.168.73.1:4726/wd/hub', desired_caps)
				time.sleep(10)
				element=driver.find_element_by_xpath("//android.widget.CheckedTextView[@text='CAIPA']")
				element.click()
				return driver, "PASS"
		elif browser=="Android":
			if user=="user1":
				desired_caps = {}
				desired_caps['platformName'] = 'Android'
				desired_caps['platformVersion'] = '6.0'
				desired_caps['deviceName'] = 'Moto'
				desired_caps['app'] = '/Users/macmini/dev/CureatrAndroidWorkSpace/APK/CureatrPlay.apk'
				driver = webdriver.Remote('http://192.168.73.1:4725/wd/hub', desired_caps)
				time.sleep(10)
				element=driver.find_element_by_xpath("//android.widget.CheckedTextView[@text='DaVita']")
				element.click()
				return driver, "PASS"
			else:
				desired_caps = {}
				desired_caps['platformName'] = 'Android'
				desired_caps['platformVersion'] = '4.4'
				desired_caps['deviceName'] = 'Moto'
				desired_caps['app'] = '/Users/macmini/dev/CureatrAndroidWorkSpace/APK/CureatrPlay.apk'
				driver = webdriver.Remote('http://192.168.73.1:4726/wd/hub', desired_caps)
				time.sleep(10)
				element=driver.find_element_by_xpath("//android.widget.CheckedTextView[@text='CAIPA']")
				element.click()
				return driver, "PASS"
	except Exception as err:
		print err
		logger.info("Exception @ LaunchWebBrowser"+str(err))
		return driver, "FAIL"

def Type(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(21):
		try:
			print "Type", getattr(Config, str(target))
			element=driver.find_element_by_xpath(getattr(Config, str(target)))
			element.send_keys(str(data))
			return "PASS", ""
		except Exception as err:
			if i>=20:
				logger.info("Exception @ Type"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				time.sleep(1)
				continue

def verifyText(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(21):
		try:
			element = driver.find_element_by_xpath(getattr(Config, str(target))).text
			element = element.encode('ascii', 'ignore').decode('ascii')
			if str(element).lower()==str(data).lower():
				return "PASS", ""
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
		except Exception as err:
			time.sleep(1)
			if i>=20:
				print err
				logger.info("Exception @ verifyText"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				continue

def Click(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(4):
		try:
			element=driver.find_element_by_xpath(getattr(Config, target))
			element.click()
			return "PASS", ""
		except Exception as err:
			if i>=3:
				logger.info("Exception @ Click"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				time.sleep(2)
				continue

def CloseBrowser(driver):
	try:
		driver.close()
		driver.quit()
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ CloseBrowser"+str(err))
		return "FAIL", ""

def CloseWebApp(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		driver.reset()
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ CloseWebApp"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		ScreenShotDirPath=subdirectory+"ScreenShots/"
		if not os.path.exists(ScreenShotDirPath):
			os.makedirs(ScreenShotDirPath)
		driver.get_screenshot_as_file(ScreenShotDirPath+TCID+"-"+TSID+"-"+DSID+'.jpg')
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ ScreenShot, "+str(err))
		return "Fail", ""
"""
def OpenWebApp(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if driver.current_url != CureatrPlayURL:
			driver.get(getattr(Config, str(target)))
			driver.implicitly_wait(5)
			Status=PageRefresh(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			if Status=="PASS":
				return "PASS", ""
			else:
				return "FAIL", ""
		else:
			return "PASS", ""
	except Exception as err:
		logger.info("Exception @ OpenWebApp"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""


def TypeText(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		element.send_keys(str(data))
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ TypeText"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def AttachFile(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if browser=="FF" or browser=="Chrome":
			data=FileAttachmentsDir+str(data)
		else:
			data=IEFileAttachmentsDir+"\\"+str(data)
		if browser=="FF":
			if str(target)=="ChangePhoto":
				element=driver.find_element_by_xpath(getattr(Config, str(target)))
				driver.execute_script("document.querySelector('[id^=change-photo-input]').parentNode.parentNode.className=''")
				driver.find_element_by_xpath(getattr(Config, str(target))).send_keys(data)
				driver.execute_script("document.querySelector('[id^=change-photo-input]').parentNode.parentNode.className='profile-button-photo'")
			else:
				element=driver.find_element_by_xpath(getattr(Config, str(target)))
				driver.execute_script("document.querySelector('[id^=fileuploadview]').parentNode.parentNode.className=''")
				driver.find_element_by_xpath(getattr(Config, str(target))).send_keys(data)
				driver.execute_script("document.querySelector('[id^=fileuploadview]').parentNode.parentNode.className='upload-manager'")
		else:
			driver.find_element_by_xpath(getattr(Config, str(target))).send_keys(data)
		if browser=="IE":
			time.sleep(10)
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ AttachFile"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def TypeGroupName(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		Text=element.get_attribute("value")
		if Text!="":
			element.clear()
		data="Group Test "+str(BackEndDrivers.random_digits(10))
		element.send_keys(str(data))
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ TypeGroupName"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def verifysearch(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		BEUserList=[]
		AppUserList=[]
		TotalList=BackEndDrivers.qa_search3(str(data), currentTestDataSheet, dataset)
		UsersList=TotalList[0]
		Status="PASS"
		for UsersLength in range(0, len(UsersList)):
			SingleUser=UsersList[UsersLength]
			print str(SingleUser.status)
			print "available" in str(SingleUser.status)
			UserFirstLastName=str(SingleUser.first_name).lower()+" "+str(SingleUser.last_name).lower()
			if str(SingleUser.status) != "None":
				if (TCID=="BusyCoverageSearch" or TCID=="OffDutyCoverageSearch" or TCID=="EditProfileCoverageSearch") and  "available" not in str(SingleUser.status):
					print "not execurted"	
				else:
					print "executed"
					BEUserList.append(UserFirstLastName)
			else:
				BEUserList.append(UserFirstLastName)
			UserFirstLastName=str(SingleUser.first_name).lower()+" "+str(SingleUser.last_name).lower()
			BEUserList.append(UserFirstLastName)
			if str(data).lower() in str(SingleUser.first_name).lower() or str(data).lower() in str(SingleUser.last_name).lower():
				Status="PASS"
			else:
				Status="FAIL"
				break
		GroupsList=TotalList[1]
		for GroupLength in range(0, len(GroupsList)):
			SingleGroup=GroupsList[GroupLength]
			BEUserList.append(str(SingleGroup['name']).lower())
			if str(data).lower() in str(SingleGroup['name']).lower():
				Status="PASS"
			else:
				Status="FAIL"
				break
		ServicesList=[]
		if TCID=="ContactsSearch" or TCID=="ToFieldSearch" or TCID=="ConversationSearch":
			ServicesList=TotalList[2]
			for ServiceLength in range(0, len(ServicesList)):
				SingleService=ServicesList[ServiceLength]
				BEUserList.append(str(SingleService.description).lower())
				if str(data).lower() in str(SingleService.description).lower():
					Status="PASS"
				else:
					Status="FAIL"
					break
		BESearchLength=len(UsersList)+len(GroupsList)+len(ServicesList)
		element=driver.find_element_by_xpath(getattr(Config, target)[4])
		Text=element.get_attribute("value")
		if Text!="":
			element.clear()
		element.send_keys(str(data))
		time.sleep(5)
		List=driver.find_elements_by_xpath(getattr(Config, target)[0])
		if BESearchLength==0 and len(List)==2 and str(TCID)=="ContactsSearch":
			ListLength=len(List)-2
			FeSearchStatus="PASS"	
		elif str(TCID)=="ContactsSearch":
			ListLength=len(List)-1
			for ListCount in range(1, len(List)+getattr(Config, target)[5]):
				name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
				title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
				if str(data).lower() in str(name.encode('ascii', 'ignore').decode('ascii')).lower() or str(data).lower() in str(title.encode('ascii', 'ignore').decode('ascii')).lower():
					FeSearchStatus="PASS"
					if " (" in name:
						AppUserList.append(str(name.split(" (")[0:][0]).lower())
					else:
						AppUserList.append(str(name).lower())
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
		else:
			try:
				ListLength=len(List)
				for ListCount in range(1, len(List)+getattr(Config, target)[5]):
					name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
					title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
					if str(data).lower() in str(name.encode('ascii', 'ignore').decode('ascii')).lower() or str(data).lower() in str(title.encode('ascii', 'ignore').decode('ascii')).lower():
						FeSearchStatus="PASS"
						if " (" in name:
							AppUserList.append(str(name.split(" (")[0:][0]).lower())
						else:
							AppUserList.append(str(name).lower())
					else:
						ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
						return "FAIL", ""
				if len(List)==0 and BESearchLength==0:
					element1=driver.find_element_by_xpath(getattr(Config, target)[6]).text
					if "No contacts found" in element1 or "There were no contacts found for" in element1 or element1=="No results are available. Please check your search." or "No patients found" in element1 or "There were no patients found for" in element1 or element1=="No results found" or element1=="No results are available. Please check your search.":
						FeSearchStatus="PASS"
					else:
						ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
						return "FAIL", ""
			except Exception as err:
				ListLength=len(List)-1
				element1=driver.find_element_by_xpath(getattr(Config, target)[6]).text
				if "No contacts found" in element1 or "There were no contacts found for" in element1 or element1=="No results are available. Please check your search." or element1=="No patients found" or element1=="No results found" or element1=="No results are available. Please check your search.":
					FeSearchStatus="PASS"
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
		CompareAPPBERestults=cmp(BEUserList.sort(),AppUserList.sort())
		if BESearchLength==ListLength and Status=="PASS" and FeSearchStatus=="PASS" and CompareAPPBERestults==0:
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ Verify Search"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def verifyToFieldsearch(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:

		TotalList=BackEndDrivers.qa_search3(str(data), currentTestDataSheet, dataset)
		UsersList=TotalList[0]
		for UsersLength in range(0, len(UsersList)):
			SingleUser=UsersList[UsersLength]
			if str(data).lower() in str(SingleUser.first_name).lower() or str(data).lower() in str(SingleUser.last_name).lower():
				Status="PASS"
			else:
				Status="FAIL"

		GroupsList=TotalList[1]
		for GroupLength in range(0, len(GroupsList)):
			SingleGroup=GroupsList[GroupLength]
			if str(data).lower() in str(SingleGroup['name']).lower():
				Status="PASS"
			else:
				Status="FAIL"

		BESearchLength=len(UsersList)+len(GroupsList)
		element=driver.find_element_by_xpath(getattr(Config, target)[4])
		Text=element.get_attribute("value")
		if Text!="":
			element.clear()
		element.send_keys(str(data))
		time.sleep(5)
		List=driver.find_elements_by_xpath(getattr(Config, target)[0])
		FeSearchStatus=Searchcontacts(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		if BESearchLength==len(List) and Status=="PASS" and FeSearchStatus[0]=="PASS":
			return "PASS", ""
		else:
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyToFieldsearch"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def verifyRecentContacts(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		time.sleep(3)
		List = driver.find_elements_by_xpath(getattr(Config, target)[0])
		for ListCount in range(1, len(List)+getattr(Config, target)[5]):
			name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
			title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
			if str(data).lower() in str(name.encode('ascii', 'ignore').decode('ascii')).lower() or str(data).lower() in str(title.encode('ascii', 'ignore').decode('ascii')).lower():
				return "PASS", ""
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyRecentContacts: "+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def TypeEMAILID(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if data =="":
			browser=""
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		Text=element.get_attribute("value")
		if Text!="":	
			element.clear()
		element.send_keys(str(browser)+str(data))
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ TypeEMAILID"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def TypePASSWORD(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if browser=="FF":
			data=getCellValueBySheet(currentTestDataSheet, dataset, "FFPASSWORD")
		elif browser=="Chrome":
			data=getCellValueBySheet(currentTestDataSheet, dataset, "CHROMEPASSWORD")
		elif browser=="IE":
			data=getCellValueBySheet(currentTestDataSheet, dataset, "IE")

		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		Text=element.get_attribute("value")
		if Text!="":	
			element.clear()
		element.send_keys(str(data))
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ TypePASSWORD"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def TypeRecepient(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if data =="":
			browser=""
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		Text=element.get_attribute("value")
		if Text!="":	
			element.clear()
		element.send_keys(str(browser)+str(data))
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ TypeRecepient"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def Maximize(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		driver.maximize_window()
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ Maximize"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def NotClickable(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		ui.WebDriverWait(driver, 5).until(EC.element_to_be_clickable((By.XPATH, getattr(Config, target))))
		element=driver.find_element_by_xpath(getattr(Config, target))
		element.click()
		if element.is_enabled():
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ NotClickable"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def DoubleClick(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(10):
		try:
			ui.WebDriverWait(driver, 1).until(EC.element_to_be_clickable((By.XPATH, getattr(Config, target))))
			element=driver.find_element_by_xpath(getattr(Config, target))
			element.click()
			time.sleep(0.300)
			element.click()
			return "PASS", ""
		except Exception as err:
			time.sleep(1)
			if i>=9:
				logger.info("Exception @ Click"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				continue
def ClickHidden(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(4):
		try:
			element=driver.find_element_by_xpath(getattr(Config, target))
			driver.execute_script("arguments[0].click();", element)
			return "PASS", ""
		except Exception as err:
			time.sleep(2)
			if i>=3:
				logger.info("Exception @ ClickHidden"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				continue

def HoverClick(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data,currentTestDataSheet, dataset,user):
	for i in range(4):
		try:
			element=driver.find_element_by_xpath(getattr(Config, str(target)))
			builder = ActionChains(driver)
			builder.move_to_element(element).perform()
			return "PASS", ""
		except Exception as err:
			time.sleep(2)
			if i>=3:
				logger.info("Exception @ HoverClick"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				continue
def CloseViewOld(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):#not using this function
	try:
		time.sleep(10)
		driver.execute_script("document.evaluate('html/body/div[8]/div/div/a', document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue")
		element = driver.find_element_by_xpath("html/body/div[8]/div/div/a")
		""
		driver.execute_script("arguments[0].click();", element)
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ HoverClick"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def ClickCss(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		a=driver.find_element_by_css_selector(getattr(Config, str(target)))
		a.click()
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ ClickCss"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def CloseWebApp(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(4):
		try:
			driver.delete_all_cookies()
			driver.get(getattr(Config, str("CureatrPlayURL")))
			ui.WebDriverWait(driver, 60).until(EC.element_to_be_clickable((By.XPATH, getattr(Config, str("ChangeOrg")))))
			a=driver.find_element_by_xpath(getattr(Config, str("ChangeOrg")))
			a.click()
			return "PASS", ""
		except Exception as err:
			if i>=3:
				try:
					Status=PageRefresh(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					ui.WebDriverWait(driver, 5).until(EC.element_to_be_clickable((By.XPATH, getattr(Config, str("ChangeOrg")))))
					a=driver.find_element_by_xpath(getattr(Config, str("ChangeOrg")))
					a.click()
					if Status=="PASS":
						return "PASS", ""
					else:
						return "FAIL", ""
				except:
					logger.info("Exception @ CloseWebApp"+str(err))
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
			else:
				continue

		
def verifyAppTitle(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if data== driver.title:
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyAppTitle"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def NewMessageTitle(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		for count in range(101):
			if data== driver.title:
				return "PASS", ""
			elif data!=driver.title and count<100:
				time.sleep(.333)
				continue
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ NewMessageTitle"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def isElementVisible(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		ui.WebDriverWait(driver, 15).until(EC.visibility_of_element_located((By.XPATH, getattr(Config, str(target)))))
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ isElementVisible"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "Fail", ""

def PageRefresh(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(4): 
		try:
			ui.WebDriverWait(driver, 10).until(EC.visibility_of_element_located((By.XPATH, getattr(Config, str("WelcomeMsg")))))
			return "PASS"
		except TimeoutException:
			driver.refresh()
			if i>=3:
				logger.info("Exception @ PageRefresh"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL"
			else:
				continue

def RefreshBrowser(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(4): 
		try:
			driver.refresh()
			ui.WebDriverWait(driver, 30).until(EC.visibility_of_element_located((By.XPATH, getattr(Config, str("Logo")))))
			driver.find_element_by_xpath(getattr(Config, str("Logo"))).click()
			return "PASS", ""
		except TimeoutException:
			if i>=3:
				logger.info("Exception @ RefreshBrowser"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				continue
def navigateBack(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		driver.back()
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ navigateBack"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "Fail", ""
def LinkState(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		ClassValue=element.get_attribute("class")
		if str(data)=="enabled":
			if "disabled" in ClassValue:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				return "PASS", ""
		else:
			if "disabled" in ClassValue or "inactive" in ClassValue:
				return "PASS", ""
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ LinkState"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		


def verifyUserStatus(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		Text=element.get_attribute("class")
		if str(Text)==str(data):
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyUserStatus"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def verifyTextcss(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element = driver.find_element_by_css_selector(getattr(Config, str(target))).text
		if str(data) in element.encode('ascii', 'ignore').decode('ascii'):
			return "PASS", ""
		else:
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyTextcss"+str(err))
		return "FAIL", ""
def verifyErrorMsg(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if data!="":
			ui.WebDriverWait(driver, 15).until(EC.visibility_of_element_located((By.XPATH, getattr(Config, str(target)))))
			element = driver.find_element_by_xpath(getattr(Config, str(target))).text
			if data in str(element):
				return "PASS", ""
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
		else:
			return "PASS", ""
	except Exception as err:
		logger.info("Exception @ verifyErrorMsg"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def verifyTextContains(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(6):
		try:
			element = driver.find_element_by_xpath(getattr(Config, str(target))).text
			element = element.encode('ascii', 'ignore').decode('ascii')
			if str(data).lower() in str(element).lower():
				return "PASS", ""
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
		except Exception as err:
			if i>=5:
				logger.info("Exception @ verifyTextContains"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				time.sleep(1)
				continue

def verifyOrgText(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element = driver.find_element_by_xpath(getattr(Config, str(target))).text
		StaticText="Sign in for "
		if element==StaticText+data:
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyOrgText"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def verifyReadStatus(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		List=driver.find_elements_by_xpath(getattr(Config, target)[0])
		n=len(List)
 		for i in range(31):
			element = driver.find_element_by_xpath(getattr(Config, target)[1]+str(n)+getattr(Config, target)[2]).text
			if element[:4]==data:
				return "PASS", ""
			elif data=="Unread" and element==data:
				return "PASS", ""
			elif i<30 and element[:4]!=data:
				time.sleep(1)
				continue
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyReadStatus"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def SearchInstitution(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		List=driver.find_elements_by_xpath("//*[@id='frame']/div/div[2]/div[2]/div/div/div/div[1]/div[2]/ul/li")
		for ListCount in range(1, len(List)+1):
			ActOrg=driver.find_element_by_xpath("//*[@id='frame']/div/div[2]/div[2]/div/div/div/div[1]/div[2]/ul/li["+str(ListCount)+"]/div/div").text
			if str(ActOrg) in data:
				Status="PASS"
				break
					
		if Status=="PASS":
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ SearchInstitution"+str(err))
		return "FAIL", ""
		
def selectInstitution(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		List=driver.find_elements_by_xpath("//*[@id='frame']/div/div[2]/div[2]/div/div/div/div[1]/div[2]/ul/li")
		for ListCount in range(1, len(List)+1):
			Status=""
			ActOrg=driver.find_element_by_xpath("//*[@id='frame']/div/div[2]/div[2]/div/div/div/div[1]/div[2]/ul/li["+str(ListCount)+"]/div/div").text
			if str(ActOrg)==data:
				driver.find_element_by_xpath("//*[@id='frame']/div/div[2]/div[2]/div/div/div/div[1]/div[2]/ul/li["+str(ListCount)+"]/div/div").click()
				Status="PASS"
				return "PASS", ""
		
		if Status=="":
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ selectInstitution"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def selectQuickMsg(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		List=driver.find_elements_by_xpath(getattr(Config, target)[0])
		for ListCount in range(1, len(List)+1):
			Status=""
			ActOrg=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
			if str(ActOrg)==data:
				driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).click()
				Status="PASS"
				return "PASS", ""
		if Status=="":
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ selectQuickMsg"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def DeleteAllQuickMsgs(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		List=driver.find_elements_by_xpath(getattr(Config, target)[0])
		ListCount=len(List)
		while ListCount>0:
			Status=""
			ActOrg=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
			if str(ActOrg)=="Delete":
				time.sleep(2)
				driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).click()
				driver.find_element_by_xpath(getattr(Config, target)[3]).click()
				time.sleep(2)
				Status=verifyQuickMsgList(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				ListCount=ListCount-1
		if Status[0]=="PASS":
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ DeleteAllQuickMsgs"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def ClearText(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		element.clear()
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ ClearText"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def verifyTextBoxValue(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		TextBoxValue=element.get_attribute("value")
		if str(data) in str(TextBoxValue):
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyTextBoxValue"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def isEncrypted(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		FieldType=element.get_attribute("type")
		if FieldType=="password" and str(data).lower()=="true":
			return "PASS", ""
		elif FieldType=="text" and str(data).lower()=="false":
			return "PASS", ""
		elif FieldType=="textarea" and str(data).lower()=="false":
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ isEncrypted"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def isTextBoxWrapped(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		FieldType=element.get_attribute("type")
		GetTagName=element.tag_name
		if (FieldType=="password" or FieldType=="text" or FieldType=="textarea") and (GetTagName=="input" or GetTagName=="textarea"):
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ isTextBoxWrapped"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def verifyMuteIconVissibility(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		FieldType=element.get_attribute("class")
		if str(data)=="YES" and "mute-icon" in str(FieldType):
			return "PASS", ""
		elif str(data)=="NO" and "mute-icon" not in str(FieldType):
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyMuteIconVissibility"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def verifyWaterMarkVailability(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		FieldType=element.get_attribute("placeholder")
		GetFieldValue=element.get_attribute("value")
		if data==FieldType+GetFieldValue:
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyWaterMarkVailability"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def verifyWaterMarkUnVailability(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		FieldType=element.get_attribute("placeholder")
		GetFieldValue=element.get_attribute("value")
		if data==FieldType+GetFieldValue:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
		else:
			return "PASS", ""
	except Exception as err:
		logger.info("Exception @ verifyWaterMarkUnVailability"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def verifySignIn(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if Correct_Data=="Y":
			element1 = driver.find_element_by_xpath(getattr(Config, "InboxLink")).is_displayed()
			element2 = driver.find_element_by_xpath(getattr(Config, "ContactsLink")).is_displayed()
			if element1==True and element2==True:
				return "PASS", "YES"
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", "NO"
		else:
			element1 = driver.find_element_by_xpath(getattr(Config, "WelcomeMsg")).is_displayed()
			element2 = driver.find_element_by_xpath(getattr(Config, "GetSupportLink")).is_displayed()
			if element1==True and element2==True:
				return "PASS", "NO"
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", "NO"
	except Exception as err:
		logger.info("Exception @ verifySignIn"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", "NO"

def verifyCPDailog(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		for i in range(30):
			driver.find_element_by_xpath(getattr(Config, "SaveChangesLink")).is_displayed()
			time.sleep(1)
			if i>=29:
				driver.refresh()
				ui.WebDriverWait(driver, 30).until(EC.visibility_of_element_located((By.XPATH, getattr(Config, str("SaveChangesLink")))))
				data="Test1234"
				target="NewPasswordTextBox"
				Status1=Typecharbychar(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				target="ReTypePasswordTextBox"
				Status2=Typecharbychar(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				driver.find_element_by_xpath(getattr(Config, "SaveChangesLink")).click()
				return "PASS", ""
			else:	
				continue
	except Exception as err:
		return "PASS", ""

def verifySignOut(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element1 = driver.find_element_by_xpath(getattr(Config, "WelcomeMsg")).is_displayed()
		element2 = driver.find_element_by_xpath(getattr(Config, "GetSupportLink")).is_displayed()
		if element1==True and element2==True:
			return "PASS", "YES"
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", "NO"
	except Exception as err:
		logger.info("Exception @ verifySignOut"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", "NO"
		
def verifyFirstTimeSignIn(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if Correct_Data=="Y":
			element1 = driver.find_element_by_xpath(getattr(Config, "NewPasswordLabel")).is_displayed()
			element2 = driver.find_element_by_xpath(getattr(Config, "ReTypePasswordLabel")).is_displayed()
			if element1==True and element2==True:
				return "PASS", "YES"
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", "NO"
		else:
			time.sleep(5)
			element1 = driver.find_element_by_xpath(getattr(Config, "WelcomeMsg")).is_displayed()
			element2 = driver.find_element_by_xpath(getattr(Config, "GetSupportLink")).is_displayed()
			if element1==True and element2==True:
				return "PASS", "NO"
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", "NO"
	except Exception as err:
		logger.info("Exception @ verifyFirstTimeSignIn"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", "NO"
	
def verifyChangePassword(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if data=="Y":
			element1 = driver.find_element_by_xpath(getattr(Config, target)[0]).is_displayed()
			element2 = driver.find_element_by_xpath(getattr(Config, target)[1]).is_displayed()
			if element1==True and element2==True:
				return "PASS", "YES"
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", "NO"
		else:
			time.sleep(5)
			element1 = driver.find_element_by_xpath(getattr(Config, target)[2]).is_displayed()
			element2 = driver.find_element_by_xpath(getattr(Config, target)[3]).is_displayed()
			if element1==True and element2==True:
				return "PASS", "NO"
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", "NO"
	except Exception as err:
		logger.info("Exception @ verifyChangePassword"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", "NO"

def wait(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		time.sleep(data)
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ wait"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def DriverWait(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		WebDriverWait(driver, 10).until(wait_for_text_to_start_with(By.find_element_by_xpath, getattr(Config, str(target)), data))
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ DriverWait"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def ImageComparision(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element = driver.find_element_by_xpath(getattr(Config, target))
		BasePath=element.get_attribute("src")
		#driver.get(BasePath)
		driver.execute_script("$(window.open('"+BasePath+"'))")
		if browser=="FF":
			driver.switch_to_window(driver.window_handles[len(driver.window_handles)-1])
		else:
			driver.switch_to_window(driver.window_handles[1])
		driver.get_screenshot_as_file(ImagesDir+browser+"ActFile.png")
		if browser=="Chrome":
			driver.execute_script("window.top.close();")
		time.sleep(1)
		driver.switch_to_window(driver.window_handles[0])
		i1 = Image.open(ImagesDir+browser+"ActFile.png")
		i2 = Image.open(ImagesDir+str(browser)+data, 'r')
		#assert i1.mode == i2.mode, "Different kinds of images."
		#assert i1.size == i2.size, "Different sizes."
		pairs = izip(i1.getdata(), i2.getdata())
		
		if len(i1.getbands()) == 1:
			# for gray-scale jpegs
			dif = sum(abs(p1-p2) for p1,p2 in pairs)
		else:
			dif = sum(abs(c1-c2) for p1,p2 in pairs for c1,c2 in zip(p1,p2))

		ncomponents = i1.size[0] * i1.size[1] * 3
		if (dif / 255.0 * 100)/ncomponents==0.0 or (dif / 255.0 * 100)/ncomponents<=0.06:
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ ImageComparision"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def ImageFilesComparision(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		for count in range (1, 100):
			if os.path.isfile(OthersDir+browser+"Downloads/"+data):
				break
			else:
				time.sleep(1)
		#file1 = open(OthersDir+browser+"Downloads/"+data, 'r')
		#file2 = open(AttachmentsDir+data, 'r')
		#f1_line = file1.readline()
		#f2_line = file2.readline()
		i1 = Image.open(OthersDir+browser+"Downloads/"+data, 'r')
		i2 = Image.open(AttachmentsDir+data, 'r')
		#assert i1.mode == i2.mode, "Different kinds of images."
		#assert i1.size == i2.size, "Different sizes."
		pairs = izip(i1.getdata(), i2.getdata())
		if len(i1.getbands()) == 1:
			# for gray-scale jpegs
			dif = sum(abs(p1-p2) for p1,p2 in pairs)
		else:
			dif = sum(abs(c1-c2) for p1,p2 in pairs for c1,c2 in zip(p1,p2))

		ncomponents = i1.size[0] * i1.size[1] * 3
		if (dif / 255.0 * 100)/ncomponents==0.0:
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ ImageFilesComparision"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def PdfComparision(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		oldfile=[]
		newfile=[]
		for line in get_pdf_content_lines(ExpPdfFile):
			oldfile.append(str(line))
			
		for line in get_pdf_content_lines(ActPdfFile):
			newfile.append(str(line.encode('ascii', 'ignore').decode('ascii')))
			
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ PdfComparision"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def get_pdf_content_lines(pdf_file_path):
	
	with open(pdf_file_path) as f:
		pdf_reader = PdfFileReader(f)
		for page in pdf_reader.pages:
			for line in page.extractText().splitlines():
				yield line

def ExcelComparision(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for count in range (1, 100):
			if os.path.isfile(OthersDir+browser+"Downloads/"+data):
				break
			else:
				time.sleep(1)
	wb = openpyxl.load_workbook(OthersDir+browser+"Downloads/"+data, 'r')
	wb1 = openpyxl.load_workbook(AttachmentsDir+data, 'r')
	sheet1 = wb.get_sheet_by_name('Suite')
	sheet2 = wb1.get_sheet_by_name('Suite')
	if sheet1.max_row ==sheet2.max_row:
		if  sheet1.max_column==sheet2.max_column:
			for i in range(1, sheet1.max_column):
				for j in range(1,sheet1.max_row):
					if sheet1.cell(row = j, column = i).value != sheet2.cell(row = j, column = i).value:
						return "FAIL", ""
					else:
						Status="PASS"
	if Status=="PASS":
		return "PASS", ""
	else:
		return "FAIL", ""

def VoiceComparision(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		audiodiff.equal('airplane.flac', 'airplane.m4a')
		audiodiff.audio_equal('airplane.flac', 'airplane.m4a')
		audiodiff.tags_equal('airplane.flac', 'airplane.m4a')
	except Exception as err:
		logger.info("Exception @ ImageComparision"+str(err))
		return "FAIL", ""

def DocxComparision(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for count in range (1, 100):
		if os.path.isfile(OthersDir+browser+"Downloads/"+data):
			break
		else:
			time.sleep(1)
	document=Document(OthersDir+browser+"Downloads/"+data)
	document2=Document(AttachmentsDir+data)
	for p,p1 in zip(document.paragraphs,document2.paragraphs):
		if p.text != p1.text:
			return "FAIL", ""
		else:
			return "PASS", ""

def VerifyProfileImage(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		FieldType=element.get_attribute("src")
		if "user" in FieldType or "profile" in FieldType:
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def verifyTermsofService(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		#driver.switch_to.frame("tos_update_text")
		TSelement = driver.find_element_by_xpath(getattr(Config, str(target))).text
		TSelement = TSelement.encode('ascii', 'ignore').decode('ascii')
		ActFilePath=TOSFileDir+'ActFile.txt'
		OpenActFile = open(ActFilePath, 'w')
		OpenActFile.write(str(TSelement))
		OpenActFile.close()
		Status=verifyTextFiles(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		if Status=="PASS":
			return "PASS", ""
		else:
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyTermsofService"+str(err))
		return "FAIL", ""

def verifyTextFiles(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		#file1 = open(TOSFileDir+'ActFile.txt', 'r')
		#file2 = open(TOSFileDir+'ExpFile.txt', 'r')
		for count in range (1, 100):
			if os.path.isfile(OthersDir+browser+"Downloads/"+'CureatrTextFile.txt'):
				break
			else:
				time.sleep(1)
		file1 = open(OthersDir+browser+"Downloads/"+'CureatrTextFile.txt', 'r')
		file2 = open(AttachmentsDir+'CureatrTextFile.txt', 'r')
		f1_line = file1.readline()
		f2_line = file2.readline()
		while f1_line != '' or f2_line != '':
			f1_line = f1_line.rstrip()
			f2_line = f2_line.rstrip()
			if f1_line==f2_line:
				Status="PASS"
			else:
				Status="FAIL"
				break
			f1_line = file1.readline()
			f2_line = file2.readline()
		
		if Status=="PASS":
			return "PASS", ""
		else:
			return "FAIl", ""
		file1.close()
		file2.close()
	except Exception as err:
		logger.info("Exception @ verifyTextFiles"+str(err))
		return "FAIL", ""

def DeleteFiles(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		folder=OthersDir+browser+"Downloads"
		filelist = [ f for f in os.listdir(folder) if f.endswith(data) ]
		for f in filelist:
			os.remove(folder+"/"+f)
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ DeleteFiles"+str(err))
		return "FAIL", ""

def verifyPatientSearch(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		BEPatientsList=[]
		FEPatientsList=[]
		PatientsList=BackEndDrivers.patient_search(str(data), currentTestDataSheet, dataset)
		Status="PASS"
		for PatientsLength in range(0, len(PatientsList)):
			SinglePatient=PatientsList[PatientsLength]
			PatientName=str(SinglePatient['name']).lower()
			pid=str(SinglePatient['pid']).lower()
			sex=str(SinglePatient['sex']).lower()
			dob=str(SinglePatient['dob']).lower()
			PatientDetails=PatientName+","+pid+","+sex+","+dob
			BEPatientsList.append(PatientDetails)
			if str(data).lower() in str(SinglePatient['name']).lower() or str(data).lower() in str(SinglePatient['pid']).lower():
				Status="PASS"
			else:
				Status="FAIL"
				break
		time.sleep(5)
		List = driver.find_elements_by_xpath(getattr(Config, target)[0])
		Status=""
		for ListCount in range(1, len(List)+getattr(Config, target)[5]):
			Appname=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
			AppMRN=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
			AppDOB=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[7]).text
			AppDOB=AppDOB.encode('ascii', 'ignore').decode('ascii')
			Appname=str(Appname).lower()
			Apppid=str(AppMRN.split(": ")[0:][1]).lower()
			AppSex=AppDOB[:4].lower()
			AppDOB=AppDOB[-25:].lower()
			AppPatientDetails=str(Appname)+","+str(Apppid)+","+str(AppSex)+","+str(AppDOB)
			FEPatientsList.append(AppPatientDetails)
			if str(data).lower() in str(Appname.encode('ascii', 'ignore').decode('ascii')).lower() or str(data).lower() in str(Apppid.encode('ascii', 'ignore').decode('ascii')).lower() or str(data).lower() in str(AppDOB.encode('ascii', 'ignore').decode('ascii')).lower():
				Status="PASS"
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
		if Status=="PASS":
			return "PASS", ""
		return "PASS", ""

	except Exception as err:
		logger.info("Exception @ verifyPatientSearch"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def verifyPatientInfo(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		BEPatientsList=[]
		FEPatientsList=[]
		PatientsList=BackEndDrivers.patient_search(str(data), currentTestDataSheet, dataset)
		Status="PASS"
		for PatientsLength in range(0, len(PatientsList)):
			SinglePatient=PatientsList[PatientsLength]
			PatientName=str(SinglePatient['name']).lower()
			pid=str(SinglePatient['pid']).lower()
			sex=str(SinglePatient['sex']).lower()
			dob=str(SinglePatient['dob']).lower()
			PatientDetails=PatientName+","+pid+","+sex+","+dob
			BEPatientsList.append(PatientDetails)
			if str(data).lower() in str(SinglePatient['name']).lower() or str(data).lower() in str(SinglePatient['pid']).lower():
				Status="PASS"
			else:
				Status="FAIL"
				break
		
		Status=""
		Appname=driver.find_element_by_xpath(Config.PatientInfoName).text
		Apppid=driver.find_element_by_xpath(Config.PatienntInfoMRN).text
		AppDOB=driver.find_element_by_xpath(Config.PatientInfoDOB).text
		AppSex=driver.find_element_by_xpath(Config.PatientInfoSex).text
		AppDOB=AppDOB.encode('ascii', 'ignore').decode('ascii')
			
		AppPatientDetails=str(Appname).lower()+","+str(Apppid).lower()+","+str(AppSex).lower()+","+str(AppDOB).lower()
		FEPatientsList.append(AppPatientDetails)
		if str(data).lower() in str(Appname.encode('ascii', 'ignore').decode('ascii')).lower() or str(data).lower() in str(Apppid.encode('ascii', 'ignore').decode('ascii')).lower() or str(data).lower() in str(AppDOB.encode('ascii', 'ignore').decode('ascii')).lower():
			Status="PASS"
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
		
		CompareAPPBERestults=cmp(BEPatientsList.sort(),FEPatientsList.sort())
		if Status=="PASS" and CompareAPPBERestults==0:
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyPatientInfo"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def Searchcontacts(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		time.sleep(5)
		List = driver.find_elements_by_xpath(getattr(Config, target)[0])
		Status=""
		for ListCount in range(1, len(List)+getattr(Config, target)[5]):
			name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
			title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
			if str(data).lower() in str(name.encode('ascii', 'ignore').decode('ascii')).lower() or str(data).lower() in str(title.encode('ascii', 'ignore').decode('ascii')).lower():
				Status="PASS"
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
		
		if Status=="PASS":
			return "PASS", ""
	except Exception as err:
		logger.info("Exception @ Search Contacts Method:: "+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def selectContact(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		time.sleep(3)
		#ui.WebDriverWait(driver, 15).until(EC.visibility_of_element_located((By.XPATH, getattr(Config, target)[0])))
		#ui.WebDriverWait(driver, 15).until(EC.element_to_be_clickable((By.XPATH, getattr(Config, target)[0])))
		List = driver.find_elements_by_xpath(getattr(Config, target)[0])
		for ListCount in range(1, len(List)+getattr(Config, target)[5]):
			Status=""
			name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
			title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
			if str(data).lower() in name.lower() or str(data).lower() in title.lower():
				element=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2])
				driver.execute_script("arguments[0].click();", element)
				Status="PASS"
				return "PASS", ""
		
		if Status=="":
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ selectInstitution"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def isVissible(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(6):
		try:
			element=driver.find_element_by_xpath(getattr(Config, str(target)))
			if element.is_displayed():
				return "PASS", ""
			else:
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", "" 
		except Exception as err:
			if i>=5:
				logger.info("Exception @ isVissible"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				time.sleep(1)
				continue 

def isUploadSucessVissible(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		List=driver.find_elements_by_xpath(getattr(Config, target)[0])
		ui.WebDriverWait(driver, 30).until(EC.visibility_of_element_located((By.XPATH, getattr(Config, target)[1]+str(len(List))+getattr(Config, target)[2])))
		element=driver.find_element_by_xpath(getattr(Config, target)[1]+str(len(List))+getattr(Config, target)[2])
		if element.is_displayed():
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", "" 
	except Exception as err:
		logger.info("Exception @ isUploadCancelBtnVissible"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", "" 		

def isNotVissible(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		if element.is_displayed():
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", "" 
		else:
			return "PASS", ""
	except Exception as err:
		if "Unable to locate element:" in str(err):
			return "PASS", ""
		else:
			logger.info("Exception @ isNotVissible"+str(err))
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""  

def isEnabled(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		if element.is_enabled():
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", "" 
	except Exception as err:
		logger.info("Exception @ isEnabled"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""  
  
def delLastchars(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, target)[4])
		element.clear()
		element.send_keys(data)
		GetFieldValue=element.get_attribute("value")
		while len(GetFieldValue) !=0:
			try:
				time.sleep(3)
				List = driver.find_elements_by_xpath(getattr(Config, target)[0])
				for ListCount in range(1, len(List)+getattr(Config, target)[5]):
					name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
					title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
					if GetFieldValue.lower()  in name.lower()  or GetFieldValue.lower()  in title.lower():
						Status="PASS"
					else:
						ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
						return "FAIL", ""
				element.send_keys(Keys.BACKSPACE)
				GetFieldValue=element.get_attribute("value")
			except Exception as err:
				element1=driver.find_element_by_xpath(getattr(Config, target)[6]).text
				if "No contacts found" in element1 or "There were no contacts found for" in element1 or element1=="No results are available. Please check your search." or "No patients found" in element1 or "There were no patients found for" in element1 or element1=="No results found":
					element.send_keys(Keys.BACKSPACE)
					GetFieldValue=element.get_attribute("value")
					Status="PASS"
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
					
		if Status=="PASS":
			return "PASS", ""
	except Exception as err:
		logger.info("Exception @ delLastchars"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
		
def AddCharByChar(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		for Count, Char in enumerate(data):
			try:
				element=driver.find_element_by_xpath(getattr(Config, target)[4])
				element.send_keys(str(Char))
				time.sleep(3)
				List = driver.find_elements_by_xpath(getattr(Config, target)[0])
				for ListCount in range(1, len(List)+getattr(Config, target)[5]):
					name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
					title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
					if str(Char).lower()  in name.lower()  or str(Char).lower()  in title.lower():
						Status="PASS"
					else:
						ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
						return "FAIL", ""
			except Exception as err:
				element1=driver.find_element_by_xpath(getattr(Config, target)[6]).text
				if "No contacts found" in element1 or "There were no contacts found for" in element1 or element1=="No results are available. Please check your search." or "No patients found" in element1 or "There were no patients found for" in element1 or element1=="No results found":
					continue
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL",""
		if Status=="PASS":
			return "PASS", ""
	except Exception as err:
		logger.info("Exception @ AddCharByChar"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def AddCharByCharManageGroup(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		for Count, Char in enumerate(data):
			element=driver.find_element_by_xpath(getattr(Config, target)[4])
			element.send_keys(str(Char))
			time.sleep(3)
			List = driver.find_elements_by_xpath(getattr(Config, target)[0])
			if len(List)==0:
				element1=driver.find_element_by_xpath(getattr(Config, target)[6]).text
				if element1=="No results are available. Please check your search.":
					Status="PASS"
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
			for ListCount in range(1, len(List)+getattr(Config, target)[5]):
				name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
				title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
				if str(Char).lower()  in name.lower()  or str(Char).lower()  in title.lower():
					Status="PASS"
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
		if Status=="PASS":
			return "PASS", ""
	except Exception as err:
		logger.info("Exception @ AddCharByCharManageGroup"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def delLastcharsManageGroup(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, target)[4])
		element.clear()
		element.send_keys(data)
		GetFieldValue=element.get_attribute("value")
		time.sleep(3)
		while len(GetFieldValue) !=0:
			List = driver.find_elements_by_xpath(getattr(Config, target)[0])
			if len(List)==0:
				element1=driver.find_element_by_xpath(getattr(Config, target)[6]).text
				if element1=="No results are available. Please check your search.":
					Status="PASS"
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
			time.sleep(1)
			for ListCount in range(1, len(List)+getattr(Config, target)[5]):
				name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
				title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
				if GetFieldValue.lower()  in name.lower()  or GetFieldValue.lower()  in title.lower():
					Status="PASS"
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
			element.send_keys(Keys.BACKSPACE)
			GetFieldValue=element.get_attribute("value")
					
		if Status=="PASS":
			return "PASS", ""
	except Exception as err:
		logger.info("Exception @ delLastcharsManageGroup"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def delLastcharsCoverageSearch(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, target)[4])
		element.clear()
		element.send_keys(data)
		GetFieldValue=element.get_attribute("value")
		print "GetFieldValue",GetFieldValue
		Status==""
		while len(GetFieldValue) !=0:
			try:
				while len(GetFieldValue) !=0:
					time.sleep(2)
					List = driver.find_elements_by_xpath(getattr(Config, target)[0])
					for ListCount in range(1, len(List)+getattr(Config, target)[5]):
						name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
						title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
						if GetFieldValue.lower()  in name.lower()  or GetFieldValue.lower()  in title.lower():
							Status="PASS"
						else:
							ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
							return "FAIL", ""
					
					element.send_keys(Keys.BACKSPACE)
					GetFieldValue=element.get_attribute("value")
			except Exception as err:
				element1=driver.find_element_by_xpath(getattr(Config, target)[6]).text
				if element1=="No results are available. Please check your search.":
					element.send_keys(Keys.BACKSPACE)
					GetFieldValue=element.get_attribute("value")
					continue
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
					
		if Status=="PASS":
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def AddCharByCharCoverageSearch(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		print "before for loop"
		element=driver.find_element_by_xpath(getattr(Config, target)[4])
		element.ClearText
		for Count, Char in enumerate(data):
			print "inside for loop"
			try:
				print "try starting"
				element=driver.find_element_by_xpath(getattr(Config, target)[4])
				element.send_keys(str(Char))
				time.sleep(2)
				List = driver.find_elements_by_xpath(getattr(Config, target)[0])
				for ListCount in range(1, len(List)+getattr(Config, target)[5]):
					name=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
					title=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
					if str(Char).lower()  in name.lower()  or str(Char).lower()  in title.lower():
						Status="PASS"
					else:
						ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
						return "FAIL", ""
			except Exception as err:
				element1=driver.find_element_by_xpath(getattr(Config, target)[6]).text
				#time.sleep(5)
				if element1=="No results are available. Please check your search.":
					continue
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL",""
		if Status=="PASS":
			return "PASS", ""
	except Exception as err:
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""
def MaximizeComposeScreen(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		ui.WebDriverWait(driver, 15).until(EC.visibility_of_element_located((By.XPATH, getattr(Config, target))))
		time.sleep(1)
		element = driver.find_element_by_xpath(getattr(Config, str(target))).size
		if (str(element)=="{'width': 482, 'height': 283}" or str(element)=="{'width': 482, 'height': 563}" or str(element)=="{'width': 482, 'height': 331}") and str(data)=="Maximize":
			return "PASS", ""
		elif str(element)=="{'width': 270, 'height': 41}" and str(data)=="Minimize":
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ MaximizeComposeScreen"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def verifyQuickMsgList(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		List = driver.find_elements_by_xpath("//div/div[1]/div/div[2]/ul[2]/li")
		n=len(List)
		num=4
		Status=""
		while n > 0:
			text=driver.find_element_by_xpath("//div/div[1]/div/div[2]/ul[2]/li["+str(n)+"]/div[2]/div[2]").text
			deletestate=driver.find_element_by_xpath("//div/div[1]/div/div[2]/ul[2]/li["+str(n)+"]/div[2]/a").text
			number=driver.find_element_by_xpath("//div/div[1]/div/div[2]/ul[2]/li["+str(n)+"]/div[1]/div").text
			icon=driver.find_element_by_xpath("//div/div[1]/div/div[2]/ul[2]/li["+str(n)+"]/div[2]/div[1]").is_displayed()
			if len(List)>6 and num<10:
				if str(num)==str(number) and icon==True and text is not None and str(deletestate)=="Delete":
					Status="PASS"
					num=num+1
					n=n-1
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
			elif len(List)<7:
				if str(num)==str(number) and icon==True and text is not None and str(deletestate)=="Delete":
					Status="PASS"
					num=num+1
					n=n-1
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
			else:
				if str(number)=="" and icon==True and text is not None and str(deletestate)=="Delete":
					Status="PASS"
					n=n-1
				else:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
		if Status=="PASS" and len(List)>0:
			return "PASS", ""
		elif Status=="" and len(List)==0:
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyQuickMsgList"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def verifyLatestMessage(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
 	try:
 		time.sleep(3)
 		i=1
 		while i < 30:
 			List=driver.find_elements_by_xpath(getattr(Config, target)[0])
 			n=len(List)
 			element=driver.find_element_by_xpath(getattr(Config, target)[1]+str(n)+getattr(Config, target)[2]).text
 			element = element.encode('ascii', 'ignore').decode('ascii')
 			if str(element)==str(data):
 				return "PASS", ""
 			else:
 				time.sleep(1)
 				i=i+1
 				Status="FAIL"
 		if Status=="FAIL":
 			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
 			return "FAIL", ""
 	except Exception as err:
 		logger.info("Exception @ verifyLatestMessage"+str(err))
 		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
 		return "FAIL", ""

def verifyLatestImage(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
 	try:
 		time.sleep(3)
 		timesec=1
 		while timesec < 60:
 			List=driver.find_elements_by_xpath(getattr(Config, target)[0])
 			element=driver.find_element_by_xpath(getattr(Config, target)[1]+str(len(List))+getattr(Config, target)[2]).text
 			element = element.encode('ascii', 'ignore').decode('ascii')
 			if element==data:
 				Status="PASS"
 				break
 			else:
 				Status="FAIL"
 				time.sleep(1)
 				timesec=timesec+1

 		element=driver.find_element_by_xpath(getattr(Config, target)[1]+str(len(List))+getattr(Config, target)[3]).is_displayed()
 		if element==True and Status=="PASS":
 			return "PASS", ""
 		else:
 			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
 			return "FAIL", ""
 	except Exception as err:
 		logger.info("Exception @ verifyLatestImage"+str(err))
 		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
 		return "FAIL", ""

def SelectLatestImage(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
 	try:
 		List=driver.find_elements_by_xpath(getattr(Config, target)[0])
 		element=driver.find_element_by_xpath(getattr(Config, target)[1]+str(len(List))+getattr(Config, target)[3]).is_displayed()
 		if element==True:
 			driver.find_element_by_xpath(getattr(Config, target)[1]+str(len(List))+getattr(Config, target)[3]).click()
 			return "PASS", ""
 		else:
 			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
 			return "FAIL", ""
 	except Exception as err:
 		logger.info("Exception @ SelectLatestImage"+str(err))
 		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
 		return "FAIL", ""

def Archieve(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		List=driver.find_elements_by_xpath("//*[@id='thread-collection']/li")
		Status=""
		for ListCount in range(1, len(List)+1):
			element=driver.find_element_by_xpath("//*[@id='thread-collection']/li["+str(ListCount)+"]")
			Style=element.get_attribute("style")
			if Style!="display: none;":
				driver.find_element_by_xpath("//*[@id='thread-collection']/li["+str(ListCount)+"]/div[2]/div[2]/div[1]/div[1]").click()
				return "PASS", ""
		if Status=="":
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ Archieve"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def GetProfilesImage(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(60):
		try:
			element=driver.find_element_by_xpath(getattr(Config, str(target)))
			FieldType=element.get_attribute("src")
			if "profileimages" in str(FieldType):
				addCellValueToBuff(currentTestDataSheet, dataset, "SRC", str(FieldType))
				return "PASS", ""
			else:
				time.sleep(1)
				continue
		except Exception as err:
			logger.info("Exception @ GetProfilesImage"+str(err))
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""

def VerifyProfilesImage(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		FieldType=element.get_attribute("src")
		if str(data)=="UserAvatar" and ("prompt.png" in str(FieldType) or "user.png" in str(FieldType)):
			return "PASS", ""
		elif str(data)=="ServiceAvatar" and ("avatar_role.png" in str(FieldType)):
			return "PASS", ""
		elif str(data)=='Available' and ("status_available.png" in str(FieldType)):
			return "PASS", ""
		elif str(data)=='Busy' and ("status_busy.png" in str(FieldType)):
			return "PASS", ""
		elif str(data)=='Off Duty' and ("status_off_duty.png" in str(FieldType)):
			return "PASS", ""
		elif str(data) == str(FieldType):
			addCellValueToBuff(currentTestDataSheet, dataset, "SRC", str(FieldType))
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ VerifyProfilesImage"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def Checkboxselection(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		n = 5
		while n > 0:
			element2=driver.find_element_by_xpath("//*[@id='patient_profile_modal']/div[2]/div/div[2]/div[2]/div/ul/li["+str(n)+"]/input").click()
			n = n-1
			if n==1:
				while n < 6:
					element2=driver.find_element_by_xpath("//*[@id='patient_profile_modal']/div[2]/div/div[2]/div[2]/div/ul/li["+str(n)+"]/input").click()
					time.sleep(1)
					n=n+1
				break
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ Checkboxselection"+str(err))
		return "FAIL", ""

def SelectCheckbox(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		Text=element.get_attribute("class")
		if ("checked" in str(Text) and str(data)=="uncheck"):
			driver.execute_script("arguments[0].click();", element)
			return "PASS", ""
		elif str(data)=="checked" or str(data)=="uncheck":
			driver.execute_script("arguments[0].click();", element)
			return "PASS", ""
		elif ("checked" in str(Text) and str(data)=="check") or ("checked" not in str(Text) and str(data)=="uncheck"):
			return "PASS", ""
		else:
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ SelectCheckbox"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def MultipleUsersReadStaus(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		ExpUserStatusIcon=data.split("&")[0]
		ExpFirstLastName=data.split("&")[1].lower()
		ExpReadIcon=data.split("&")[2].lower()
		List = driver.find_elements_by_xpath(getattr(Config, target)[0])
		for ListCount in range(1, len(List)+1):
			Status=""
			UserStatusIcon=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).get_attribute("src")
			FirstLastName=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[3]).text
			ReadIcon=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[5]).get_attribute("src")
			if ExpUserStatusIcon.lower() in UserStatusIcon.lower() and ExpFirstLastName.lower() in FirstLastName.lower() and ExpReadIcon.lower() in ReadIcon.lower():
				Status="PASS"
				return "PASS", ""
		
		if Status=="":
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ MultipleUsersReadStaus"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def clearcharacter(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset,user):
 	try:
  		element=driver.find_element_by_xpath(getattr(Config, str(target)))
  		FieldType=element.get_attribute("value")
  		while len(FieldType) !=0:
   			element.send_keys(Keys.BACKSPACE)
   			FieldType=element.get_attribute("value")
  		return "PASS", "" 
 	except Exception as err:
 		logger.info("Exception @ clearcharacter"+str(err))
  		return "FAIL", ""		

def Urgenttoast(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target))).click()
		alertmsg=driver.find_element_by_xpath("//div[@class='urgent-mode-enabled-alert']").text
		if str(data)==str(alertmsg):
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", "" 
	except Exception as err:
		logger.info("Exception @ Urgenttoast"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def Typecharbychar(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		for Count, Char in enumerate(data):
			element=driver.find_element_by_xpath(getattr(Config, str(target)))
			element.send_keys(str(Char))
			time.sleep(0.1)
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ Typecharbychar"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def ControlEnd(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		if browser=='Chrome':
			element=driver.find_element_by_xpath(getattr(Config, str(target)))
			element.send_keys(Keys.END)
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ ControlEnd"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def PressEnterKey(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		element.send_keys(Keys.ENTER)
		time.sleep(0.5)
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ PressEnterKey"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def switchtoframe(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		driver.switch_to_frame(driver.find_element_by_xpath(getattr(Config, str(target))))
		return "PASS", ""	
	except Exception as err:
		logger.info("Exception @ switchtoframe"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""		
		
def switchtowindow(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		driver.switch_to_window(driver.window_handles[0])
		return "PASS", ""	
	except Exception as err:
		logger.info("Exception @ switchtowindow"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""	
		
def SwitchTo(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		driver.switch_to_window(driver.window_handles[1])
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ SwitchTo"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def GroupSort(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		GroupNames=[]
		List=driver.find_elements_by_xpath(getattr(Config, target)[0])
		n=len(List)
		for ListCount in range(1, len(List)+1):
			element=driver.find_element_by_xpath(getattr(Config, target)[1]+str(ListCount)+getattr(Config, target)[2]).text
			GroupNames.append(str(element))
		CompareAPPBERestults=cmp(GroupNames,GroupNames.sort())
		return "PASS", ""
	except Exception as err:
		logger.info("Exception @ GroupSort"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def Checkboxpresent(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		Text=element.get_attribute("class")
		if ("check" in str(Text)):
			return "PASS", ""
		else:
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ Checkboxpresent"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def VerifyElementColor(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		FieldType=element.value_of_css_property('background-color')
		if str(data) == str(FieldType):
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ VerifyElementColor"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def verifyTextBoxValue2(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		element=driver.find_element_by_xpath(getattr(Config, str(target)))
		TextBoxValue=element.get_attribute("value")
		if str(data) == str(TextBoxValue):
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ verifyTextBoxValue2"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""

def listlength(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	try:
		List = driver.find_elements_by_xpath(getattr(Config, str(target)))
		if str(data)==str(len(List)):
			return "PASS", ""
		else:
			ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
			return "FAIL", ""
	except Exception as err:
		logger.info("Exception @ listlength"+str(err))
		ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
		return "FAIL", ""


def WaitForExpectedText(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user):
	for i in range(60):
		try:
			element = driver.find_element_by_xpath(getattr(Config, str(target))).text
			element = element.encode('ascii', 'ignore').decode('ascii')
			if str(data).lower() in str(element).lower():
				return "PASS", ""
			else:
				if i>=59:
					ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
					return "FAIL", ""
				else:
					time.sleep(1)
					continue
		except Exception as err:
			if i>=9:
				logger.info("Exception @ WaitForExpectedText"+str(err))
				ScreenShot(browser, driver, target, data, subdirectory, TCID, TSID, DSID, Correct_Data, currentTestDataSheet, dataset, user)
				return "FAIL", ""
			else:
				time.sleep(1)
				continue
"""